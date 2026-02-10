"""
Módulo de Exportación - PDF, Word y Excel
"""
import streamlit as st
import io
from datetime import datetime
from database import get_connection
from auth import get_current_user, require_role
from utils import section_header, msg_success, msg_error


def render():
    """Renderiza el módulo de exportación."""
    require_role(["auditor", "supervisor", "auditor_campo"])
    user = get_current_user()

    section_header("Exportación de Reportes", "Exporte proyectos y hallazgos en múltiples formatos", "📤")

    conn = get_connection()
    planes = conn.execute("SELECT * FROM planes ORDER BY anio DESC").fetchall()
    proyectos_ua = conn.execute("SELECT * FROM proyectos ORDER BY codigo_auditoria").fetchall()

    tab1, tab2, tab3 = st.tabs(["📊 Excel", "📄 Word", "📋 PDF"])

    with tab1:
        _export_excel(conn, planes, proyectos_ua)

    with tab2:
        _export_word(conn, planes)

    with tab3:
        _export_pdf(conn, planes)

    conn.close()


def _export_excel(conn, planes, proyectos_ua):
    """Exportación a Excel."""
    st.markdown("##### Exportar a Excel")

    export_type = st.selectbox("¿Qué desea exportar?", [
        "Universo Auditable",
        "Plan de Auditoría con Hallazgos",
        "Evaluación del Universo",
        "Hallazgos (todos)",
    ], key="excel_type")

    if st.button("📊 Generar Excel", type="primary", key="gen_excel"):
        try:
            import pandas as pd

            output = io.BytesIO()

            if export_type == "Universo Auditable":
                data = conn.execute("""
                    SELECT p.codigo_auditoria as "Código", p.nombre_auditoria as "Nombre",
                        p.tipo_auditoria as "Tipo", p.proceso as "Proceso",
                        p.estado as "Estado",
                        p.fecha_inicial_planificada as "Fecha Inicio Plan.",
                        p.fecha_final_planificada as "Fecha Fin Plan.",
                        p.fecha_inicio_real as "Fecha Inicio Real",
                        p.fecha_final_real as "Fecha Fin Real",
                        u1.nombre_completo as "Supervisor",
                        u2.nombre_completo as "Auditor Campo"
                    FROM proyectos p
                    LEFT JOIN usuarios u1 ON p.supervisor_id = u1.id
                    LEFT JOIN usuarios u2 ON p.auditor_campo_id = u2.id
                    ORDER BY p.codigo_auditoria
                """).fetchall()
                df = pd.DataFrame([dict(d) for d in data])
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Universo Auditable', index=False)

            elif export_type == "Plan de Auditoría con Hallazgos":
                plan_options = {p["id"]: f"{p['nombre_plan']} ({p['anio']})" for p in planes}
                if not plan_options:
                    st.warning("No hay planes disponibles.")
                    return

                # Get all plan data
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    for plan in planes:
                        # Projects
                        projs = conn.execute("""
                            SELECT pp.codigo_auditoria as "Código", pp.nombre_auditoria as "Nombre",
                                pp.tipo_auditoria as "Tipo", pp.estado as "Estado",
                                pp.fecha_inicial_planificada as "Fecha Inicio",
                                pp.fecha_final_planificada as "Fecha Fin"
                            FROM plan_proyectos pp
                            WHERE pp.plan_id = ?
                        """, (plan["id"],)).fetchall()
                        df_proj = pd.DataFrame([dict(p) for p in projs])
                        sheet_name = f"Plan {plan['anio']}"[:31]
                        if not df_proj.empty:
                            df_proj.to_excel(writer, sheet_name=sheet_name, index=False)

                    # All hallazgos
                    hallazgos = conn.execute("""
                        SELECT h.codigo_hallazgo as "Código", h.condicion as "Condición",
                            h.causa as "Causa", h.efecto as "Efecto",
                            h.recomendacion as "Recomendación", h.criterio as "Criterio",
                            h.probabilidad as "Probabilidad", h.impacto as "Impacto",
                            h.nivel_riesgo as "Nivel Riesgo", h.area as "Área",
                            h.estado as "Estado", h.fecha_asignacion as "Fecha Asignación",
                            h.fecha_compromiso as "Fecha Compromiso",
                            h.fecha_respuesta as "Fecha Respuesta",
                            pp.nombre_auditoria as "Proyecto",
                            p.nombre_plan as "Plan",
                            u.nombre_completo as "Responsable"
                        FROM hallazgos h
                        LEFT JOIN plan_proyectos pp ON h.plan_proyecto_id = pp.id
                        LEFT JOIN planes p ON h.plan_id = p.id
                        LEFT JOIN usuarios u ON h.responsable_id = u.id
                        ORDER BY h.codigo_hallazgo
                    """).fetchall()
                    df_hall = pd.DataFrame([dict(h) for h in hallazgos])
                    if not df_hall.empty:
                        df_hall.to_excel(writer, sheet_name='Hallazgos', index=False)

            elif export_type == "Evaluación del Universo":
                data = conn.execute("""
                    SELECT p.codigo_auditoria as "Código", p.nombre_auditoria as "Nombre",
                        COALESCE(e.nivel_riesgo, 0) as "Nivel Riesgo",
                        COALESCE(e.meses_ultima_auditoria, 0) as "Meses Últ. Auditoría",
                        COALESCE(e.hallazgos_ult_auditoria, 0) as "Hallazgos Últ. Aud.",
                        COALESCE(e.hallazgos_solucionados, 0) as "Hallazgos Solucionados",
                        COALESCE(e.estado_auditoria, 'N/A') as "Estado Auditoría",
                        e.fecha_auditoria as "Fecha Auditoría",
                        COALESCE(e.ciclo_rotacion, 12) as "Ciclo Rotación",
                        COALESCE(e.nivel_criticidad, 0) as "Nivel Criticidad"
                    FROM proyectos p
                    LEFT JOIN evaluacion_universo e ON p.id = e.proyecto_id
                    ORDER BY COALESCE(e.nivel_criticidad, 0) DESC
                """).fetchall()
                df = pd.DataFrame([dict(d) for d in data])
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Evaluación', index=False)

            else:  # Hallazgos todos
                data = conn.execute("""
                    SELECT h.codigo_hallazgo as "Código", h.condicion as "Condición",
                        h.causa as "Causa", h.efecto as "Efecto",
                        h.recomendacion as "Recomendación", h.criterio as "Criterio",
                        h.probabilidad as "Probabilidad", h.impacto as "Impacto",
                        h.nivel_riesgo as "Nivel Riesgo", h.area as "Área",
                        h.estado as "Estado", h.respuesta_auditado as "Respuesta Auditado",
                        h.fecha_asignacion as "F. Asignación",
                        h.fecha_compromiso as "F. Compromiso",
                        h.fecha_respuesta as "F. Respuesta",
                        u.nombre_completo as "Responsable",
                        pp.nombre_auditoria as "Proyecto",
                        p.nombre_plan as "Plan"
                    FROM hallazgos h
                    LEFT JOIN plan_proyectos pp ON h.plan_proyecto_id = pp.id
                    LEFT JOIN planes p ON h.plan_id = p.id
                    LEFT JOIN usuarios u ON h.responsable_id = u.id
                    ORDER BY h.codigo_hallazgo
                """).fetchall()
                df = pd.DataFrame([dict(d) for d in data])
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Hallazgos', index=False)

            output.seek(0)
            filename = f"Reporte_{export_type.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
            st.download_button("⬇️ Descargar Excel", data=output, file_name=filename,
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            msg_success()

        except ImportError as e:
            msg_error("Se requiere openpyxl")
        except Exception as e:
            msg_error()


def _export_word(conn, planes):
    """Exportación a Word."""
    st.markdown("##### Exportar a Word")

    plan_options = {p["id"]: f"{p['nombre_plan']} ({p['anio']})" for p in planes}
    if not plan_options:
        st.info("No hay planes disponibles.")
        return

    selected_plan = st.selectbox("Plan", list(plan_options.keys()),
                                  format_func=lambda x: plan_options[x], key="word_plan")

    if st.button("📄 Generar Word", type="primary", key="gen_word"):
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import os
            from config import ASSETS_DIR

            doc = Document()

            # Style
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(10)

            # Header with logo
            logo_path = os.path.join(ASSETS_DIR, "logo.png")
            if os.path.exists(logo_path):
                doc.add_picture(logo_path, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Title
            plan_data = dict(conn.execute("SELECT * FROM planes WHERE id = ?", (selected_plan,)).fetchone())
            title = doc.add_heading(f"Informe de Auditoría", level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph(f"Plan: {plan_data['nombre_plan']}")
            doc.add_paragraph(f"Año: {plan_data['anio']}")
            doc.add_paragraph(f"Objetivo: {plan_data.get('objetivo', 'N/A')}")
            doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

            doc.add_heading("Proyectos del Plan", level=1)

            proyectos = conn.execute("""
                SELECT * FROM plan_proyectos WHERE plan_id = ? ORDER BY codigo_auditoria
            """, (selected_plan,)).fetchall()

            for proj in proyectos:
                proj = dict(proj)
                doc.add_heading(f"{proj['codigo_auditoria']} — {proj['nombre_auditoria']}", level=2)
                doc.add_paragraph(f"Estado: {proj['estado']}")
                doc.add_paragraph(f"Tipo: {proj.get('tipo_auditoria', 'N/A')}")
                doc.add_paragraph(f"Período planificado: {proj.get('fecha_inicial_planificada', 'N/A')} a {proj.get('fecha_final_planificada', 'N/A')}")

                # Hallazgos del proyecto
                hallazgos = conn.execute("""
                    SELECT h.*, u.nombre_completo as responsable
                    FROM hallazgos h
                    LEFT JOIN usuarios u ON h.responsable_id = u.id
                    WHERE h.plan_proyecto_id = ?
                    ORDER BY h.codigo_hallazgo
                """, (proj["id"],)).fetchall()

                if hallazgos:
                    doc.add_heading("Hallazgos", level=3)
                    for h in hallazgos:
                        h = dict(h)
                        doc.add_paragraph(f"Código: {h['codigo_hallazgo']} | Riesgo: {h.get('nivel_riesgo', 'N/A')} | Estado: {h['estado']}")
                        if h.get("condicion"):
                            doc.add_paragraph(f"Condición: {h['condicion']}")
                        if h.get("recomendacion"):
                            doc.add_paragraph(f"Recomendación: {h['recomendacion']}")
                        if h.get("respuesta_auditado"):
                            doc.add_paragraph(f"Respuesta: {h['respuesta_auditado']}")
                        doc.add_paragraph("")

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            filename = f"Informe_Auditoria_{plan_data['anio']}_{datetime.now().strftime('%Y%m%d')}.docx"
            st.download_button("⬇️ Descargar Word", data=output, file_name=filename,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            msg_success()

        except ImportError:
            msg_error("Se requiere python-docx")
        except Exception as e:
            msg_error()


def _export_pdf(conn, planes):
    """Exportación a PDF."""
    st.markdown("##### Exportar a PDF")

    plan_options = {p["id"]: f"{p['nombre_plan']} ({p['anio']})" for p in planes}
    if not plan_options:
        st.info("No hay planes disponibles.")
        return

    selected_plan = st.selectbox("Plan", list(plan_options.keys()),
                                  format_func=lambda x: plan_options[x], key="pdf_plan")

    if st.button("📋 Generar PDF", type="primary", key="gen_pdf"):
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors as rl_colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
            from reportlab.lib.units import inch
            import os
            from config import ASSETS_DIR

            output = io.BytesIO()
            doc = SimpleDocTemplate(output, pagesize=letter, topMargin=50, bottomMargin=50)
            elements = []
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                          textColor=rl_colors.HexColor("#233F84"), fontSize=18)
            heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                            textColor=rl_colors.HexColor("#0D68A5"), fontSize=14)

            # Logo
            logo_path = os.path.join(ASSETS_DIR, "logo.png")
            if os.path.exists(logo_path):
                img = Image(logo_path, width=2.5*inch, height=0.7*inch)
                elements.append(img)
                elements.append(Spacer(1, 20))

            plan_data = dict(conn.execute("SELECT * FROM planes WHERE id = ?", (selected_plan,)).fetchone())

            elements.append(Paragraph("Informe de Auditoría", title_style))
            elements.append(Spacer(1, 10))
            elements.append(Paragraph(f"<b>Plan:</b> {plan_data['nombre_plan']} ({plan_data['anio']})", styles['Normal']))
            elements.append(Paragraph(f"<b>Fecha:</b> {datetime.now().strftime('%d/%m/%Y')}", styles['Normal']))
            elements.append(Spacer(1, 20))

            # Hallazgos summary table
            hallazgos = conn.execute("""
                SELECT h.codigo_hallazgo, h.condicion, h.nivel_riesgo, h.estado, h.area,
                    pp.nombre_auditoria
                FROM hallazgos h
                LEFT JOIN plan_proyectos pp ON h.plan_proyecto_id = pp.id
                WHERE h.plan_id = ?
                ORDER BY h.codigo_hallazgo
            """, (selected_plan,)).fetchall()

            elements.append(Paragraph("Resumen de Hallazgos", heading_style))
            elements.append(Spacer(1, 10))

            if hallazgos:
                table_data = [["Código", "Proyecto", "Riesgo", "Estado", "Área"]]
                for h in hallazgos:
                    h = dict(h)
                    table_data.append([
                        h["codigo_hallazgo"],
                        (h.get("nombre_auditoria") or "")[:25],
                        h.get("nivel_riesgo") or "N/A",
                        h["estado"],
                        h.get("area") or "N/A"
                    ])

                table = Table(table_data, colWidths=[1.2*inch, 1.8*inch, 1*inch, 1.2*inch, 1.3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor("#233F84")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), rl_colors.white),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#E2E8F0")),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#F5F7FA")]),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ]))
                elements.append(table)
            else:
                elements.append(Paragraph("No hay hallazgos registrados en este plan.", styles['Normal']))

            doc.build(elements)
            output.seek(0)

            filename = f"Informe_PDF_{plan_data['anio']}_{datetime.now().strftime('%Y%m%d')}.pdf"
            st.download_button("⬇️ Descargar PDF", data=output, file_name=filename,
                               mime="application/pdf")
            msg_success()

        except ImportError:
            msg_error("Se requiere reportlab")
        except Exception as e:
            msg_error()
